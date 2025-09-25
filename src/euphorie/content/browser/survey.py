from ..interfaces import IQuestionContainer
from ..module import IModule
from ..profilequestion import IProfileQuestion
from ..risk import IRisk
from ..solution import ISolution
from ..survey import get_tool_type
from ..survey import ISurvey
from ..survey import ISurveyAddSchema
from ..utils import DragDropHelper
from ..utils import IToolTypesInfo
from Acquisition import aq_chain
from Acquisition import aq_inner
from Acquisition import aq_parent
from datetime import date
from docx.api import Document
from euphorie.client.docx.compiler import _sanitize_html
from euphorie.client.docx.compiler import delete_paragraph
from euphorie.client.docx.compiler import IdentificationReportCompiler
from euphorie.client.docx.html import HtmlToWord
from euphorie.client.docx.views import IdentificationReportDocxView
from euphorie.content import MessageFactory as _
from euphorie.content.behaviors.toolcategory import IToolCategory
from euphorie.content.country import ICountry
from OFS.event import ObjectClonedEvent
from plone import api
from plone.dexterity.browser.add import DefaultAddForm
from plone.dexterity.browser.add import DefaultAddView
from plone.dexterity.browser.edit import DefaultEditForm
from plone.memoize.instance import memoize
from plonetheme.nuplone.skin import actions
from plonetheme.nuplone.utils import formatDate
from Products.CMFCore.utils import getToolByName
from Products.Five import BrowserView
from Products.Five.browser.pagetemplatefile import ViewPageTemplateFile
from Products.statusmessages.interfaces import IStatusMessage
from z3c.form.interfaces import HIDDEN_MODE
from ZODB.POSException import ConflictError
from zope.component import getMultiAdapter
from zope.component import getUtility
from zope.container.interfaces import INameChooser
from zope.event import notify
from zope.i18n import translate


class SurveyBase(BrowserView):
    def _morph(self, child):
        state = getMultiAdapter((child, self.request), name="plone_context_state")
        return {
            "id": child.id,
            "title": child.title,
            "url": state.view_url(),
            "is_profile_question": IProfileQuestion.providedBy(child),
        }

    @property
    def modules_and_profile_questions(self):
        return [
            self._morph(child)
            for child in self.context.values()
            if IModule.providedBy(child) or IProfileQuestion.providedBy(child)
        ]

    @property
    def modules(self):
        """List modules in current context."""
        return [
            self._morph(child)
            for child in self.context.values()
            if IModule.providedBy(child)
        ]

    @property
    def risks(self):
        """List risks in current context."""
        return [
            self._morph(child)
            for child in self.context.values()
            if IRisk.providedBy(child)
        ]

    @property
    @memoize
    def portal_transforms(self):
        return api.portal.get_tool("portal_transforms")

    def get_safe_html(self, text):
        if not text:
            return ""
        data = self.portal_transforms.convertTo(
            "text/x-html-safe", text, mimetype="text/html"
        )
        return data.getData()


class SurveyView(SurveyBase, DragDropHelper):
    @property
    @memoize
    def training_questions(self):
        return self.context.listFolderContents(
            {"portal_type": "euphorie.training_question"}
        )

    @property
    def show_training_questions(self):
        if not api.portal.get_registry_record(
            "euphorie.use_training_module", default=False
        ):
            return False
        return getattr(self.context, "enable_test_questions", False)

    @property
    def group(self):
        return aq_parent(aq_inner(self.context))

    @property
    def use_tool_category(self):
        return IToolCategory.providedBy(self.context)


class AddForm(DefaultAddForm):
    """Custom add form for :obj:`Survey` instances.

    This form is needlessly complicated: it should use a schema and a
    vocabulary to offer a list of template surveys, but this is
    impossible since vocabulary factories always get a None context. See
    http://code.google.com/p/dexterity/issues/detail?id=125
    """

    portal_type = "euphorie.survey"
    schema = ISurveyAddSchema
    template = ViewPageTemplateFile("templates/survey_add.pt")

    def surveys(self):
        templates = [
            {"id": survey.id, "title": survey.title}
            for survey in self.context.values()
            if ISurvey.providedBy(survey)
        ]
        return templates

    def copyTemplate(self, source, title):
        target = aq_inner(self.context)
        try:
            source._notifyOfCopyTo(target, op=0)
        except ConflictError:
            raise

        copy = source._getCopy(target)
        copy.title = title
        chooser = INameChooser(target)
        copy.id = chooser.chooseName(None, copy)
        target._setObject(copy.id, copy)

        copy = target[copy.id]  # Acquisition-wrap
        copy.wl_clearLocks()
        copy._postCopy(target, op=0)
        notify(ObjectClonedEvent(target[copy.id]))
        return copy

    def createAndAdd(self, data):
        surveygroup = aq_inner(self.context)
        template = surveygroup[self.request.form["survey"]]
        survey = self.copyTemplate(template, data["title"])
        self.immediate_view = survey.absolute_url()
        return survey


class AddView(DefaultAddView):
    form = AddForm


class EditForm(DefaultEditForm):
    def applyChanges(self, data):
        changes = super().applyChanges(data)
        if changes:
            # Reindex our parents title.
            catalog = getToolByName(self.context, "portal_catalog")
            catalog.indexObject(aq_parent(aq_inner(self.context)))
        return changes

    @property
    @memoize
    def portal_transforms(self):
        return api.portal.get_tool("portal_transforms")

    def get_safe_html(self, text):
        if not text:
            return ""
        data = self.portal_transforms.convertTo(
            "text/x-html-safe", text, mimetype="text/html"
        )
        return data.getData()

    def updateWidgets(self):
        super().updateWidgets()
        if not api.portal.get_registry_record(
            "euphorie.use_integrated_action_plan", default=False
        ):
            self.widgets["integrated_action_plan"].mode = HIDDEN_MODE
        if not api.portal.get_registry_record(
            "euphorie.use_training_module", default=False
        ):
            self.widgets["enable_web_training"].mode = HIDDEN_MODE
            self.widgets["enable_test_questions"].mode = HIDDEN_MODE
            self.widgets["num_training_questions"].mode = HIDDEN_MODE
        else:
            for obj in aq_chain(aq_inner(self.context)):
                if ICountry.providedBy(obj):
                    if not obj.enable_web_training:
                        self.widgets["enable_web_training"].mode = HIDDEN_MODE
                        self.widgets["enable_test_questions"].mode = HIDDEN_MODE
                        self.widgets["num_training_questions"].mode = HIDDEN_MODE
                        break
        for fname in ("introduction",):
            value = self.widgets[fname].value or ""
            safe_value = self.get_safe_html(value)
            if value != safe_value:
                self.widgets[fname].value = safe_value


class Delete(actions.Delete):
    """Special delete action class which prevents deletion of published surveys
    or of the last survey in a group."""

    def verify(self, container, context):
        flash = IStatusMessage(self.request).addStatusMessage
        if container.published_survey == context:
            flash(
                _(
                    "message_no_delete_published_survey",
                    default="You cannot delete an OiRA Tool version that is "
                    "published. Please unpublish it first.",
                ),
                "error",
            )
            self.request.response.redirect(context.absolute_url())
            return False

        count = 0
        for survey in container.values():
            if ISurvey.providedBy(survey):
                if count == 1:
                    # We have at least one other survey,
                    # so we can delete this one.
                    return True
                count += 1

        flash(
            _(
                "message_delete_no_last_survey",
                default="This is the only version of the OiRA Tool and can "
                "therefore not be deleted. Did you perhaps want to "
                "remove the OiRA Tool itself?",
            ),
            "error",
        )
        self.request.response.redirect(context.absolute_url())
        return False


class ContentsOfSurveyCompiler(IdentificationReportCompiler):
    def __init__(self, context, request=None):
        """Read the docx template and initialize some instance attributes that
        will be used to compile the template."""
        self.context = context
        self.request = request
        self.template = Document(self._template_filename)

        self.compiler = HtmlToWord()

        self.use_existing_measures = False
        self.tool_type = get_tool_type(self.context)
        self.tti = getUtility(IToolTypesInfo)
        self.italy_special = False

    def set_session_title_row(self, data):
        request = self.request

        # Remove existing paragraphs
        for paragraph in self.template.paragraphs:
            delete_paragraph(paragraph)

        header = self.template.sections[0].header
        header_table = header.tables[0]
        header_table.cell(0, 0).paragraphs[0].text = data["title"]
        header_table.cell(0, 1).paragraphs[0].text = formatDate(request, date.today())

        if getattr(self.context, "published"):
            footer_txt = "This OiRA tool was last published {date}.".format(
                date=self.context.published.strftime("%Y/%m/%d")
            )
        else:
            footer_txt = "This OiRA tool is currently not published."

        footer = self.template.sections[0].footer
        paragraph = footer.tables[0].cell(0, 0).paragraphs[0]
        paragraph.style = "Footer"
        paragraph.text = footer_txt

    def add_report_section(self, nodes, heading, **extra):
        doc = self.template
        doc.add_paragraph(heading, style="Heading 1")

        for node in nodes:
            title = "[{}] {}".format(
                translate(_(node.typus), target_language=self.lang), node.title
            )
            number = node.number

            doc.add_paragraph(
                f"{number} {title}", style="Heading %d" % (node.depth + 1)
            )

            if node.typus == "Risk":
                doc.add_paragraph(
                    "[%s] %s"
                    % (
                        translate(
                            _(
                                "label_problem_description",
                                default="Negative statement",
                            ),
                            target_language=self.lang,
                        ),
                        node.problem_description,
                    ),
                    style="Measure Heading",
                )

            description = node.description

            self.compiler(_sanitize_html(description or ""), doc)

            if node.typus != "Risk":
                continue

            if not extra.get("skip_legal_references", False):
                legal_reference = getattr(node, "legal_reference", None)
                if legal_reference and legal_reference.strip():
                    doc.add_paragraph()
                    legal_heading = translate(
                        _(
                            "header_legal_references",
                            default="Legal and policy references",
                        ),
                        target_language=self.lang,
                    )
                    doc.add_paragraph(legal_heading, style="Legal Heading")
                    self.compiler(_sanitize_html(legal_reference), doc)


class Node:
    title = ""
    typus = ""
    depth = 0
    number = ""
    description = ""
    legal_reference = None
    problem_description = None

    def __init__(
        self,
        title,
        typus,
        depth,
        number,
        description="",
        legal_reference=None,
        problem_description=None,
    ):
        self.title = title
        self.typus = typus
        self.depth = depth
        self.number = number
        self.description = description
        self.legal_reference = legal_reference
        self.problem_description = problem_description


class MockWebHelpers:
    can_view_session = True


class ContentsOfSurvey(IdentificationReportDocxView):
    _compiler = ContentsOfSurveyCompiler
    nodes = []
    # Webhelpers are actually not needed for _this_ computation. But the
    # parent class's __call__ method checks if the user can_view_session. In the
    # context of client, that makes sense. Here, there is no session to check for,
    # so we override that check.
    webhelpers = MockWebHelpers()

    def __init__(self, request, context):
        super().__init__(request, context)
        self.nodes = []

    def AddToTree(self, node, depth, number):
        legal_reference = None
        problem_description = None
        if IRisk.providedBy(node):
            typus = "Risk"
            legal_reference = getattr(node, "legal_reference", None)
            problem_description = getattr(node, "problem_description", None)
        elif IProfileQuestion.providedBy(node):
            typus = "Profile question"
        else:
            typus = "Module"
        self.nodes.append(
            Node(
                title=node.title,
                typus=typus,
                depth=depth,
                number=".".join(number[:depth]),
                description=node.description,
                legal_reference=legal_reference,
                problem_description=problem_description,
            )
        )

        if IQuestionContainer.providedBy(node):
            i = 0
            depth += 1
            if len(number) < depth:
                number.append("0")
            for child in node.values():
                i += 1
                number[depth - 1] = str(i)
                self.AddToTree(child, depth, number)
        elif IRisk.providedBy(node):
            i = 0
            depth += 1
            number.append("0")
            for child in node.values():
                if not ISolution.providedBy(child):
                    continue
                i += 1
                number[depth - 1] = str(i)
                description = "<ul>"
                for field in ("action", "requirements"):
                    value = getattr(child, field, "") or ""
                    if value:
                        description = f"{description}<li>{value}</li>"
                description = f"{description}</ul>"
                self.nodes.append(
                    Node(
                        title=child.description,
                        typus="Measure",
                        depth=depth,
                        number=".".join(number[:depth]),
                        description=description,
                    )
                )

    def get_survey_nodes(self):
        i = 0
        for child in self.context.values():
            if IQuestionContainer.providedBy(child):
                i += 1
                self.AddToTree(child, depth=1, number=[str(i)])
        return self.nodes

    def get_data(self, for_download=False):
        """Gets the data structure in a format suitable for `DocxCompiler`"""
        title = self.context.aq_parent.title
        data = {
            "title": title,
            "heading": "",
            "section_headings": [title],
            "nodes": [self.get_survey_nodes()],
        }
        return data

    @property
    def _filename(self):
        """Return the document filename."""
        filename = _(
            "filename_tool_contents",
            default="Contents of OIRA tool ${title}",
            mapping=dict(title=self.context.title),
        )
        return f"{translate(filename, context=self.request)}.docx"


class FindSolutionTitleDuplications(BrowserView):
    """"""

    @property
    def measures(self):
        brains = api.content.find(
            portal_type="euphorie.solution",
            path="/".join(self.context.getPhysicalPath()),
        )
        for brain in brains:
            obj = brain.getObject()
            if obj.description in obj.action:
                yield obj


class FindToolsWithDuplications(BrowserView):
    @property
    def tools(self):
        brains = api.content.find(
            portal_type="euphorie.survey",
            path="/".join(self.context.getPhysicalPath()),
        )
        tools = []
        for brain in brains:
            obj = brain.getObject()
            if not obj.published:
                continue
            view = api.content.get_view(
                name="find-solution-title-duplications",
                context=obj,
                request=self.request,
            )
            for measure in view.measures:
                tools.append(obj)
                # we're done if we found at least one measure
                break
        return tools
